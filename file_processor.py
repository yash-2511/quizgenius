import os
import logging
from typing import Optional

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text
        
    Raises:
        Exception: If extraction fails
    """
    if fitz is None:
        raise Exception("PyMuPDF not available. Please install pymupdf package.")
    
    try:
        text = ""
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            text += page_text + "\n"
        
        doc.close()
        
        # Clean up the text
        text = text.strip()
        
        if not text:
            raise Exception("No text could be extracted from the PDF")
        
        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text
        
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text
        
    Raises:
        Exception: If extraction fails
    """
    if Document is None:
        raise Exception("python-docx not available. Please install python-docx package.")
    
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = text.strip()
        
        if not text:
            raise Exception("No text could be extracted from the DOCX file")
        
        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from file based on its extension
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text
        
    Raises:
        Exception: If extraction fails or file type not supported
    """
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        logger.error(f"File processing error for {file_path}: {str(e)}")
        raise

def get_file_info(file_path: str) -> dict:
    """
    Get basic information about a file
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        dict: File information
    """
    try:
        stat = os.stat(file_path)
        return {
            'size': stat.st_size,
            'extension': os.path.splitext(file_path)[1].lower(),
            'name': os.path.basename(file_path)
        }
    except Exception as e:
        logger.error(f"Error getting file info for {file_path}: {str(e)}")
        return {}
